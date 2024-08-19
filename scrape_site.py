import os
import asyncio
import argparse
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError
from docx import Document
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup, NavigableString, Tag

# Function to create a Word document from structured HTML content
def save_as_word(title, structured_content, folder):
    doc = Document()
    doc.add_heading(title.strip(), level=1)

    for element in structured_content:
        if isinstance(element, str):
            doc.add_paragraph(element.strip(), style="Normal")
        elif isinstance(element, tuple):
            text, style = element
            if style == "heading":
                doc.add_heading(text.strip(), level=1)
            elif style == "heading_2":
                doc.add_heading(text.strip(), level=2)
            elif style == "heading_3":
                doc.add_heading(text.strip(), level=3)
            elif style == "intense quote":
                doc.add_paragraph(text.strip(), style="Quote")
            else:
                doc.add_paragraph(text.strip(), style="Normal")

    file_name = f"{title[:30].replace(' ', '_').replace('/', '-')}.docx"
    file_path = os.path.join(folder, file_name)
    doc.save(file_path)
    print(f"Saved: {file_path}")

# Function to get internal links from a page
async def get_internal_links(page, base_url):
    links = await page.query_selector_all('a')
    internal_links = set()
    for link in links:
        href = await link.get_attribute('href')
        if href and href.startswith('/'):
            href = urljoin(base_url, href)
        if href and urlparse(href).hostname == urlparse(base_url).hostname:
            internal_links.add(href)
    return internal_links

# Function to extract and structure page content using BeautifulSoup
def extract_structured_content(content_html):
    soup = BeautifulSoup(content_html, 'html.parser')

    structured_content = []

    tag_handlers = {
        'h1': 'heading',
        'h2': 'heading_2',
        'h3': 'heading_3',
        'h4': 'heading_3',
        'h5': 'heading_3',
        'h6': 'heading_3',
        'p': 'Normal',
        'li': 'Normal',
        'blockquote': 'intense quote'
    }

    for tag in soup.body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'blockquote', 'div']):
        # Handle divs with long text
        if tag.name == 'div' and len(tag.get_text(strip=True)) > 100:
            structured_content.append((tag.get_text(separator=' ', strip=True), "Normal"))
        elif isinstance(tag, Tag) and tag.name in tag_handlers:
            style = tag_handlers[tag.name]
            text = tag.get_text(separator=' ', strip=True)
            if text:
                structured_content.append((text, style))

    return structured_content

# Function to determine if the page is an HTML page using Playwright
async def is_html_page(page):
    content_type = await page.evaluate("document.contentType")
    return content_type.lower() == 'text/html'

# Preliminary check to see if the URL is likely to be HTML
def is_likely_html(url):
    non_html_extensions = ('.pdf', '.jpg', '.jpeg', '.png', '.gif', '.mp3', '.mp4', '.zip', '.exe', '.docx', '.xlsx', '.pptx')
    return not url.lower().endswith(non_html_extensions)

# Function to scrape a page using Playwright
async def scrape_page(url, folder, visited_pages, max_pages):
    if len(visited_pages) >= max_pages:
        return

    if not is_likely_html(url):
        print(f"Skipping non-HTML or unsupported file: {url}")
        return

    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()

        retries = 2
        while retries > 0:
            try:
                await page.goto(url, timeout=30000)  # 30-second timeout
                break
            except PlaywrightTimeoutError:
                retries -= 1
                print(f"Timeout occurred while loading: {url}. Retrying...")

        if retries == 0:
            print(f"Failed to load page after retries: {url}. Skipping.")
            await browser.close()
            return

        # Check if the page is HTML before proceeding
        if not await is_html_page(page):
            print(f"Skipping non-HTML content: {url}")
            await browser.close()
            return

        title = await page.title()
        content_html = await page.content()
        structured_content = extract_structured_content(content_html)

        save_as_word(title, structured_content, folder)

        base_url = urlparse(url).scheme + '://' + urlparse(url).hostname
        internal_links = await get_internal_links(page, base_url)

        visited_pages.add(url)
        await browser.close()

        return internal_links

# Function to crawl and scrape multiple pages
async def scrape_website(start_url, folder, max_pages):
    if not os.path.exists(folder):
        os.makedirs(folder)

    visited_pages = set()
    to_visit = {start_url}

    while to_visit and len(visited_pages) < max_pages:
        current_url = to_visit.pop()
        if current_url in visited_pages:
            continue

        print(f"Visiting: {current_url}")
        internal_links = await scrape_page(current_url, folder, visited_pages, max_pages)

        if internal_links:
            to_visit.update(internal_links - visited_pages)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape a website and save content as Word documents.")
    parser.add_argument("url", type=str, help="The starting URL of the website to scrape.")
    parser.add_argument("max_pages", type=int, help="The maximum number of pages to scrape.")
    parser.add_argument("--output_folder", type=str, default="docx", help="The folder where Word documents will be saved.")

    args = parser.parse_args()

    asyncio.run(scrape_website(args.url, args.output_folder, args.max_pages))
